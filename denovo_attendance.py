"""
Denovo Attendance Report Generator
Processes ZKTeco biometric clocking data into a polished Word document.
"""

import csv
import sys
import argparse
from collections import defaultdict
from datetime import datetime, date, timedelta, time
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

WORK_START = time(8, 0)
ROUND_MINUTES = 15
MERGE_WINDOW_MINUTES = 5

# Fixed break schedule (start, end) — only deducted if shift spans the window
FIXED_BREAKS = [
    (time(10, 0), time(10, 15)),   # 15 min
    (time(13, 0), time(13, 30)),   # 30 min
    (time(16, 0), time(16, 15)),   # 15 min
]

# Colours
RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xED, 0x7D, 0x31)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_BLUE = RGBColor(0x1F, 0x39, 0x64)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
HEADER_BLUE = RGBColor(0x2E, 0x74, 0xB5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WEEK_TOTAL_BG = RGBColor(0xE2, 0xEF, 0xDA)
GRAND_TOTAL_BG = RGBColor(0xFF, 0xFF, 0xCC)

# Page layout (A4 portrait, in DXA = twentieths of a point)
PAGE_WIDTH_DXA = 11906   # 210 mm
PAGE_HEIGHT_DXA = 16838  # 297 mm
MARGIN_DXA = 720         # 12.7 mm (0.5 inch)
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - 2 * MARGIN_DXA  # 10466


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def load_employees(csv_path: str) -> dict:
    """
    Returns {emp_id: full_name} from employee CSV.
    Columns: ID, ?, first_name, last_name, ...
    """
    employees = {}
    with open(csv_path, newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        for row in reader:
            if len(row) < 4:
                continue
            emp_id = row[0].strip()
            first = row[2].strip()
            last = row[3].strip()
            if emp_id:
                employees[emp_id] = f"{first} {last}".strip()
    return employees


def load_clockings(dat_path: str) -> dict:
    """
    Returns {emp_id: {date: [datetime, ...]}} from tab-separated .dat file.
    Columns: emp_id, timestamp (YYYY-MM-DD HH:MM:SS), ...
    """
    clockings: dict = defaultdict(lambda: defaultdict(list))
    with open(dat_path, newline="", encoding="utf-8-sig") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            parts = line.split("\t")
            if len(parts) < 2:
                continue
            emp_id = parts[0].strip()
            ts_str = parts[1].strip()
            try:
                ts = datetime.strptime(ts_str, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                continue
            clockings[emp_id][ts.date()].append(ts)
    return clockings


# ---------------------------------------------------------------------------
# Time arithmetic
# ---------------------------------------------------------------------------

def round_to_nearest(dt: datetime, minutes: int = ROUND_MINUTES) -> datetime:
    """Round a datetime to the nearest N minutes."""
    total_minutes = dt.hour * 60 + dt.minute
    rounded = round(total_minutes / minutes) * minutes
    new_hour = rounded // 60
    new_minute = rounded % 60
    # Handle overflow at midnight
    if new_hour >= 24:
        new_hour = 23
        new_minute = 59
    return dt.replace(hour=new_hour, minute=new_minute, second=0, microsecond=0)


def cap_to_work_start(dt: datetime) -> datetime:
    """Cap any clocking before WORK_START to WORK_START."""
    if dt.time() < WORK_START:
        return dt.replace(hour=WORK_START.hour, minute=WORK_START.minute, second=0, microsecond=0)
    return dt


def merge_duplicates(clockings: list, window_minutes: int = MERGE_WINDOW_MINUTES) -> list:
    """Merge clockings within window_minutes of each other (keep first)."""
    if not clockings:
        return []
    result = [clockings[0]]
    for c in clockings[1:]:
        if (c - result[-1]).total_seconds() / 60 >= window_minutes:
            result.append(c)
    return result


def minutes_to_hhmm(total_minutes: float) -> str:
    """Convert float minutes to H:MM string."""
    if total_minutes < 0:
        total_minutes = 0
    h = int(total_minutes) // 60
    m = int(total_minutes) % 60
    return f"{h}:{m:02d}"


def overlap_minutes(seg_start: time, seg_end: time, brk_start: time, brk_end: time) -> float:
    """Return minutes of overlap between two time ranges."""
    start = max(seg_start, brk_start)
    end = min(seg_end, brk_end)
    if end <= start:
        return 0.0
    return (datetime.combine(date.today(), end) - datetime.combine(date.today(), start)).total_seconds() / 60


def fixed_break_deduction(seg_start: time, seg_end: time) -> float:
    """Sum of fixed breaks that fall within a work segment."""
    total = 0.0
    for bs, be in FIXED_BREAKS:
        total += overlap_minutes(seg_start, seg_end, bs, be)
    return total


# ---------------------------------------------------------------------------
# Daily calculation
# ---------------------------------------------------------------------------

def process_day(raw_clockings: list) -> dict:
    """
    Process a list of raw datetime clockings for a single day.

    Returns dict:
        clockings_display: list of time strings (after rounding/capping)
        gross_minutes: float
        break_minutes: float
        net_minutes: float
        note: str
        break_is_actual: bool  (True = from clocked breaks, False = fixed schedule)
        break_over_limit: bool
    """
    # Sort, cap, round, merge
    sorted_c = sorted(raw_clockings)
    capped = [cap_to_work_start(c) for c in sorted_c]
    rounded = [round_to_nearest(c) for c in capped]
    merged = merge_duplicates(sorted(rounded))

    display = [dt.strftime("%H:%M") for dt in merged]
    n = len(merged)

    result = {
        "clockings_display": display,
        "gross_minutes": 0.0,
        "break_minutes": 0.0,
        "net_minutes": 0.0,
        "note": "",
        "break_is_actual": False,
        "break_over_limit": False,
    }

    if n == 0:
        result["note"] = "No clockings"
        return result

    if n == 1:
        result["note"] = "Single clocking - check"
        return result

    first_in = merged[0].time()
    last_out = merged[-1].time()
    gross = (merged[-1] - merged[0]).total_seconds() / 60
    result["gross_minutes"] = gross

    if n == 2:
        # Fixed break schedule
        brk = fixed_break_deduction(first_in, last_out)
        result["break_minutes"] = brk
        result["net_minutes"] = gross - brk
        result["break_is_actual"] = False

    elif n % 2 == 0:
        # 4+ clockings: pair up into work segments
        work_minutes = 0.0
        for i in range(0, n, 2):
            seg_start = merged[i].time()
            seg_end = merged[i + 1].time()
            seg_dur = (merged[i + 1] - merged[i]).total_seconds() / 60
            work_minutes += seg_dur

        # Also deduct fixed breaks from within work segments
        extra_brk = 0.0
        for i in range(0, n, 2):
            extra_brk += fixed_break_deduction(merged[i].time(), merged[i + 1].time())

        brk = gross - work_minutes + extra_brk
        result["break_minutes"] = brk
        result["net_minutes"] = gross - brk
        result["break_is_actual"] = True

    else:
        # Odd number: best-effort pairing, flag it
        result["note"] = "Odd clockings - check"
        work_minutes = 0.0
        for i in range(0, n - 1, 2):
            work_minutes += (merged[i + 1] - merged[i]).total_seconds() / 60
        brk = gross - work_minutes
        result["break_minutes"] = brk
        result["net_minutes"] = gross - brk
        result["break_is_actual"] = True

    # Flag excessive breaks (> 60 min standard allowance)
    result["break_over_limit"] = result["break_minutes"] > 60

    return result


# ---------------------------------------------------------------------------
# Date utilities
# ---------------------------------------------------------------------------

def week_start(d: date) -> date:
    """Return Monday of the week containing d."""
    return d - timedelta(days=d.weekday())


def date_range(start: date, end: date):
    """Yield each date from start to end inclusive."""
    current = start
    while current <= end:
        yield current
        current += timedelta(days=1)


def get_weeks(start: date, end: date) -> list:
    """Return list of (week_monday, [dates_in_range]) tuples."""
    weeks = defaultdict(list)
    for d in date_range(start, end):
        weeks[week_start(d)].append(d)
    return sorted(weeks.items())


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for s in tcPr.findall(qn("w:shd")):
        tcPr.remove(s)
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # Insert shd in correct position (after tcW and tcBorders, before tcMar/vAlign)
    _insert_shd(tcPr, shd)


def _insert_shd(tcPr, shd_elem):
    """Insert shd element in correct OOXML order within tcPr."""
    ORDER = [
        qn("w:cnfStyle"), qn("w:tcW"), qn("w:gridSpan"), qn("w:hMerge"),
        qn("w:vMerge"), qn("w:tcBorders"), qn("w:shd"), qn("w:noWrap"),
        qn("w:tcMar"), qn("w:textDirection"), qn("w:tcFitText"),
        qn("w:vAlign"), qn("w:hideMark"), qn("w:headers"),
    ]
    children = list(tcPr)
    # Remove old shd if any
    for c in children:
        if c.tag == qn("w:shd"):
            tcPr.remove(c)
    # Find insertion point
    insert_after = None
    for c in tcPr:
        try:
            if ORDER.index(c.tag) < ORDER.index(qn("w:shd")):
                insert_after = c
        except ValueError:
            pass
    if insert_after is not None:
        insert_after.addnext(shd_elem)
    else:
        tcPr.insert(0, shd_elem)


def set_cell_borders(cell, border_size=4, color="auto"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    # Replace existing
    for b in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(b)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is not None:
        tcW.addnext(tcBorders)
    else:
        tcPr.insert(0, tcBorders)


def set_cell_width(cell, width_dxa: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.insert(0, tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_run_color(run, rgb: RGBColor):
    rPr = run._r.get_or_add_rPr()
    color_elem = rPr.find(qn("w:color"))
    if color_elem is None:
        color_elem = OxmlElement("w:color")
        rPr.append(color_elem)
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    color_elem.set(qn("w:val"), hex_color)


def add_run(para, text, bold=False, italic=False, size_pt=10, color=None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color:
        set_run_color(run, color)
    return run


def patch_page_size(doc: Document):
    """Ensure A4 portrait with no stale w:orient attribute."""
    sectPr = doc.sections[0]._sectPr
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.append(pgSz)
    pgSz.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    pgSz.set(qn("w:h"), str(PAGE_HEIGHT_DXA))
    # Remove stale orient attribute if present
    orient_attr = qn("w:orient")
    if orient_attr in pgSz.attrib:
        del pgSz.attrib[orient_attr]


def patch_zoom(doc: Document):
    """Ensure w:zoom has required w:percent attribute."""
    settings_elem = doc.settings.element
    zoom = settings_elem.find(qn("w:zoom"))
    if zoom is not None and not zoom.get(qn("w:percent")):
        zoom.set(qn("w:percent"), "100")


def set_col_widths(table, widths_dxa: list):
    """Set explicit column widths (in DXA) for a table."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    # Also set on each cell
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                set_cell_width(cell, widths_dxa[i])


def cell_para(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Get or create the single paragraph in a cell, set alignment."""
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = alignment
    # Remove space before/after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    dxa = int(height_cm * 567)  # 1 cm = 567 DXA
    trHeight.set(qn("w:val"), str(dxa))
    trHeight.set(qn("w:hRule"), "atLeast")


def style_header_cell(cell, text, widths_hint=None, font_size=9):
    set_cell_bg(cell, HEADER_BLUE)
    set_cell_borders(cell, border_size=4, color="FFFFFF")
    p = cell_para(cell)
    add_run(p, text, bold=True, size_pt=font_size, color=WHITE)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_break(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    return para


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def build_title_page(doc: Document, company: str, period_start: date, period_end: date,
                     weeks: list, employee_totals: dict, employees: dict):
    """Build the title/summary page."""
    section = doc.sections[0]

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, company, bold=True, size_pt=22, color=DARK_BLUE)

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, "Attendance Report", bold=True, size_pt=16, color=HEADER_BLUE)

    # Period
    period_str = f"{period_start.strftime('%d %B %Y')} — {period_end.strftime('%d %B %Y')}"
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(24)
    add_run(p3, period_str, size_pt=12, color=DARK_BLUE)

    # Divider line
    _add_horizontal_line(doc)

    # Calculation rules
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(16)
    p4.paragraph_format.space_after = Pt(4)
    add_run(p4, "Calculation Rules", bold=True, size_pt=11, color=DARK_BLUE)

    rules = [
        "Work start is capped to 08:00 — any earlier clocking is treated as 08:00.",
        "All clockings are rounded to the nearest 15 minutes.",
        "Duplicate punches within 5 minutes of each other are merged.",
        "2 clockings (in/out only): a fixed break schedule is applied — 10:00–10:15 (15 min), 13:00–13:30 (30 min), 16:00–16:15 (15 min) — only deducted if the shift spans that window.",
        "4+ clockings (employee clocked breaks): clockings are paired (1st in/2nd out/3rd in/4th out…) to determine actual work segments. Fixed breaks are also deducted from within those segments.",
        "Odd number of clockings: best-effort pairing; flagged as 'Odd clockings - check'.",
        "Single clocking only: flagged as 'Single clocking - check'; hours cannot be calculated.",
        "Break value shown with * indicates actual clocked breaks were used (not fixed schedule).",
        "Break value shown in red indicates total break exceeds 60 minutes.",
        "Days with odd clockings are shown in orange.",
    ]
    for rule in rules:
        pr = doc.add_paragraph(style="List Bullet")
        pr.paragraph_format.space_before = Pt(1)
        pr.paragraph_format.space_after = Pt(1)
        add_run(pr, rule, size_pt=9)

    # Summary table
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(20)
    p5.paragraph_format.space_after = Pt(6)
    add_run(p5, "Weekly Totals Summary", bold=True, size_pt=11, color=DARK_BLUE)

    _build_summary_table(doc, weeks, employee_totals, employees)

    # Page break after title page
    add_page_break(doc)


def _add_horizontal_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{HEADER_BLUE[0]:02X}{HEADER_BLUE[1]:02X}{HEADER_BLUE[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_summary_table(doc: Document, weeks: list, employee_totals: dict, employees: dict):
    """Build the summary table with employee rows and week columns."""
    n_weeks = len(weeks)
    # Columns: Employee Name + one per week + Grand Total
    n_cols = 1 + n_weeks + 1
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Column widths
    name_w = 2200
    total_w = 1000
    remaining = CONTENT_WIDTH_DXA - name_w - total_w
    week_w = remaining // n_weeks if n_weeks else remaining
    # Adjust last week column to absorb rounding
    week_ws = [week_w] * n_weeks
    if n_weeks:
        week_ws[-1] = CONTENT_WIDTH_DXA - name_w - total_w - week_w * (n_weeks - 1)
    widths = [name_w] + week_ws + [total_w]
    set_col_widths(table, widths)

    # Header row
    hdr = table.rows[0]
    set_row_height(hdr, 0.7)
    style_header_cell(hdr.cells[0], "Employee")
    for i, (wmon, _) in enumerate(weeks):
        wend = wmon + timedelta(days=6)
        label = f"W/C {wmon.strftime('%d %b')}"
        style_header_cell(hdr.cells[1 + i], label, font_size=8)
    style_header_cell(hdr.cells[-1], "Total")

    # Employee rows
    sorted_emps = sorted(employee_totals.keys(),
                         key=lambda eid: employees.get(eid, eid))
    for eid in sorted_emps:
        row = table.add_row()
        set_row_height(row, 0.55)
        # Name cell
        p = cell_para(row.cells[0], WD_ALIGN_PARAGRAPH.LEFT)
        name = employees.get(eid, eid)
        add_run(p, f"  {name}", size_pt=9)
        set_cell_borders(row.cells[0], border_size=4, color="CCCCCC")

        grand_total = 0.0
        for i, (wmon, wdates) in enumerate(weeks):
            week_mins = sum(
                employee_totals[eid].get(d, 0.0)
                for d in wdates
            )
            grand_total += week_mins
            p = cell_para(row.cells[1 + i])
            add_run(p, minutes_to_hhmm(week_mins) if week_mins else "—", size_pt=9)
            set_cell_borders(row.cells[1 + i], border_size=4, color="CCCCCC")

        # Grand total
        p = cell_para(row.cells[-1])
        add_run(p, minutes_to_hhmm(grand_total), bold=True, size_pt=9)
        set_cell_borders(row.cells[-1], border_size=4, color="CCCCCC")
        set_cell_bg(row.cells[-1], LIGHT_GRAY)

    # Alternate row shading
    for ri, row in enumerate(table.rows[1:], start=1):
        if ri % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, LIGHT_GRAY)


# ---------------------------------------------------------------------------
# Employee detail page
# ---------------------------------------------------------------------------

def build_employee_page(doc: Document, emp_id: str, emp_name: str,
                        daily_data: dict, weeks: list,
                        period_start: date, period_end: date):
    """Build one page per employee with daily attendance table."""

    # Employee header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, emp_name, bold=True, size_pt=13, color=DARK_BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    add_run(p2, f"Employee ID: {emp_id}  |  Period: {period_start.strftime('%d %b %Y')} – {period_end.strftime('%d %b %Y')}",
            size_pt=9, color=HEADER_BLUE)

    _add_horizontal_line(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Table columns: Date | Day | Clockings | Gross | Break | Net | Note
    col_labels = ["Date", "Day", "Clockings", "Gross Hrs", "Break", "Net Hrs", "Note"]
    col_widths = [900, 700, 4766, 800, 700, 800, 1800]
    # Adjust clockings column to fill exactly
    col_widths[2] = CONTENT_WIDTH_DXA - sum(col_widths) + col_widths[2]

    table = doc.add_table(rows=1, cols=len(col_labels))
    table.style = "Table Grid"
    set_col_widths(table, col_widths)

    # Header row
    hdr = table.rows[0]
    set_row_height(hdr, 0.65)
    for i, label in enumerate(col_labels):
        style_header_cell(hdr.cells[i], label, font_size=9)

    grand_net = 0.0
    grand_gross = 0.0
    grand_break = 0.0

    for wmon, wdates in weeks:
        week_net = 0.0
        week_gross = 0.0
        week_break = 0.0

        for d in wdates:
            if d < period_start or d > period_end:
                continue
            day_result = daily_data.get(d)
            is_odd = day_result and "Odd clockings" in day_result.get("note", "")
            text_color = ORANGE if is_odd else BLACK

            row = table.add_row()
            set_row_height(row, 0.5)

            # Date
            p = cell_para(row.cells[0])
            add_run(p, d.strftime("%d %b"), size_pt=9, color=text_color)

            # Day
            p = cell_para(row.cells[1])
            add_run(p, d.strftime("%a"), size_pt=9, color=text_color)

            # Clockings
            p = cell_para(row.cells[2])
            if day_result and day_result["clockings_display"]:
                clock_str = " | ".join(day_result["clockings_display"])
            else:
                clock_str = "—"
            add_run(p, clock_str, size_pt=8, color=text_color)

            # Gross
            p = cell_para(row.cells[3])
            if day_result and day_result["gross_minutes"]:
                add_run(p, minutes_to_hhmm(day_result["gross_minutes"]), size_pt=9, color=text_color)
                week_gross += day_result["gross_minutes"]
            else:
                add_run(p, "—", size_pt=9)

            # Break
            p = cell_para(row.cells[4])
            if day_result and day_result["gross_minutes"]:
                brk_mins = day_result["break_minutes"]
                brk_str = minutes_to_hhmm(brk_mins)
                if day_result["break_is_actual"]:
                    brk_str += "*"
                brk_color = RED if day_result["break_over_limit"] else text_color
                r = add_run(p, brk_str, size_pt=9, color=brk_color)
                if day_result["break_over_limit"]:
                    r.bold = True
                week_break += brk_mins
            else:
                add_run(p, "—", size_pt=9)

            # Net
            p = cell_para(row.cells[5])
            if day_result and day_result["gross_minutes"]:
                net_mins = day_result["net_minutes"]
                add_run(p, minutes_to_hhmm(net_mins), size_pt=9, color=text_color)
                week_net += net_mins
            else:
                add_run(p, "—", size_pt=9)

            # Note
            p = cell_para(row.cells[6], WD_ALIGN_PARAGRAPH.LEFT)
            note = day_result.get("note", "") if day_result else ""
            add_run(p, note, size_pt=8, color=text_color, italic=bool(note))

            # Alternate shading for weekend
            if d.weekday() >= 5:
                for cell in row.cells:
                    set_cell_bg(cell, LIGHT_GRAY)

        # Weekly subtotal row
        week_row = table.add_row()
        set_row_height(week_row, 0.55)
        for cell in week_row.cells:
            set_cell_bg(cell, WEEK_TOTAL_BG)
        p = cell_para(week_row.cells[0])
        wend = wmon + timedelta(days=6)
        add_run(p, "W/T", bold=True, size_pt=8, color=DARK_BLUE)
        p = cell_para(week_row.cells[1])
        add_run(p, wmon.strftime("%d %b"), size_pt=8, color=DARK_BLUE)
        p = cell_para(week_row.cells[2])
        add_run(p, f"Week of {wmon.strftime('%d %b %Y')}", size_pt=8, color=DARK_BLUE, italic=True)
        p = cell_para(week_row.cells[3])
        add_run(p, minutes_to_hhmm(week_gross), bold=True, size_pt=9, color=DARK_BLUE)
        p = cell_para(week_row.cells[4])
        add_run(p, minutes_to_hhmm(week_break), bold=True, size_pt=9, color=DARK_BLUE)
        p = cell_para(week_row.cells[5])
        add_run(p, minutes_to_hhmm(week_net), bold=True, size_pt=9, color=DARK_BLUE)
        cell_para(week_row.cells[6])

        grand_net += week_net
        grand_gross += week_gross
        grand_break += week_break

    # Grand total row
    gt_row = table.add_row()
    set_row_height(gt_row, 0.6)
    for cell in gt_row.cells:
        set_cell_bg(cell, GRAND_TOTAL_BG)
    p = cell_para(gt_row.cells[0])
    add_run(p, "TOTAL", bold=True, size_pt=9, color=DARK_BLUE)
    cell_para(gt_row.cells[1])
    cell_para(gt_row.cells[2])
    p = cell_para(gt_row.cells[3])
    add_run(p, minutes_to_hhmm(grand_gross), bold=True, size_pt=9, color=DARK_BLUE)
    p = cell_para(gt_row.cells[4])
    add_run(p, minutes_to_hhmm(grand_break), bold=True, size_pt=9, color=DARK_BLUE)
    p = cell_para(gt_row.cells[5])
    add_run(p, minutes_to_hhmm(grand_net), bold=True, size_pt=10, color=DARK_BLUE)
    cell_para(gt_row.cells[6])


# ---------------------------------------------------------------------------
# Main report builder
# ---------------------------------------------------------------------------

def generate_report(dat_path: str, csv_path: str, output_path: str,
                    company: str = "Denovo Apparel Ltd",
                    period_start: date = None, period_end: date = None):

    print(f"Loading employees from {csv_path}...")
    employees = load_employees(csv_path)
    print(f"  Found {len(employees)} employees.")

    print(f"Loading clockings from {dat_path}...")
    raw_clockings = load_clockings(dat_path)
    print(f"  Found clocking data for {len(raw_clockings)} employee IDs.")

    # Determine date range from data if not specified
    all_dates = [
        d
        for emp_data in raw_clockings.values()
        for d in emp_data.keys()
    ]
    if not all_dates:
        print("ERROR: No clocking data found.")
        sys.exit(1)

    if period_start is None:
        period_start = min(all_dates)
    if period_end is None:
        period_end = max(all_dates)

    print(f"  Period: {period_start} to {period_end}")

    weeks = get_weeks(period_start, period_end)

    # Process all days for all employees
    processed: dict = {}          # {emp_id: {date: day_result}}
    employee_totals: dict = {}    # {emp_id: {date: net_minutes}}

    all_emp_ids = set(raw_clockings.keys())

    for eid in sorted(all_emp_ids):
        processed[eid] = {}
        employee_totals[eid] = {}
        for d in date_range(period_start, period_end):
            day_clocks = raw_clockings[eid].get(d, [])
            if day_clocks:
                result = process_day(day_clocks)
                processed[eid][d] = result
                employee_totals[eid][d] = result["net_minutes"]

    # Build document
    doc = Document()

    # Set default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Page setup
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.left_margin = Twips(MARGIN_DXA)
    section.right_margin = Twips(MARGIN_DXA)
    section.top_margin = Twips(MARGIN_DXA)
    section.bottom_margin = Twips(MARGIN_DXA)
    patch_page_size(doc)

    print("Building title page...")
    build_title_page(doc, company, period_start, period_end,
                     weeks, employee_totals, employees)

    # One page per employee — only those with clocking data
    emp_ids_sorted = sorted(
        all_emp_ids,
        key=lambda eid: employees.get(eid, eid)
    )

    for idx, eid in enumerate(emp_ids_sorted):
        emp_name = employees.get(eid, f"Unknown ({eid})")
        print(f"  Building page for {emp_name} ({eid})...")

        build_employee_page(
            doc, eid, emp_name,
            processed[eid], weeks,
            period_start, period_end
        )

        # Page break between employees (not after the last one)
        if idx < len(emp_ids_sorted) - 1:
            add_page_break(doc)

    patch_zoom(doc)

    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_date(s: str) -> date:
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    raise argparse.ArgumentTypeError(f"Cannot parse date: {s}")


def main():
    parser = argparse.ArgumentParser(
        description="Generate Denovo Apparel attendance report from ZKTeco biometric data."
    )
    parser.add_argument("dat_file", help="Path to ZKTeco .dat clocking file")
    parser.add_argument("employee_csv", help="Path to employee CSV file")
    parser.add_argument("output", nargs="?", default="attendance_report.docx",
                        help="Output .docx path (default: attendance_report.docx)")
    parser.add_argument("--company", default="Denovo Apparel Ltd",
                        help="Company name for report header")
    parser.add_argument("--start", type=parse_date, metavar="YYYY-MM-DD",
                        help="Report start date (default: earliest date in data)")
    parser.add_argument("--end", type=parse_date, metavar="YYYY-MM-DD",
                        help="Report end date (default: latest date in data)")
    args = parser.parse_args()

    generate_report(
        dat_path=args.dat_file,
        csv_path=args.employee_csv,
        output_path=args.output,
        company=args.company,
        period_start=args.start,
        period_end=args.end,
    )


if __name__ == "__main__":
    main()
